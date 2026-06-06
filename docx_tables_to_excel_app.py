from __future__ import annotations

import re
import threading
from datetime import date, datetime, time
from decimal import Decimal
from pathlib import Path
from queue import Queue
from typing import Callable

import pandas as pd
from docx import Document
from openpyxl import load_workbook

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk
except ImportError as error:
    class _MissingTk:
        class Tk:
            pass

    tk = _MissingTk()
    filedialog = None
    messagebox = None
    ttk = None
    TKINTER_ERROR = error
else:
    TKINTER_ERROR = None


APP_TITLE = "Magic Button: Word ↔ Excel"
MAX_SHEET_NAME_LENGTH = 31
INVALID_SHEET_CHARS = r"[\[\]\:\*\?\/\\]"


def clean_text(value: object) -> str:
    """Convert a cell value to clean text for Word and Excel output."""
    if value is None:
        return ""

    if isinstance(value, datetime):
        text = value.strftime("%d.%m.%Y %H:%M")
    elif isinstance(value, date):
        text = value.strftime("%d.%m.%Y")
    elif isinstance(value, time):
        text = value.strftime("%H:%M")
    elif isinstance(value, float) and value.is_integer():
        text = str(int(value))
    elif isinstance(value, Decimal):
        text = format(value, "f")
    else:
        text = str(value)

    return " ".join(text.split())


def make_unique_sheet_name(file_name: str, table_number: int, used_names: set[str]) -> str:
    """Build a valid unique Excel sheet name no longer than 31 characters."""
    base_name = Path(file_name).stem
    base_name = re.sub(INVALID_SHEET_CHARS, "_", base_name).strip()

    if not base_name:
        base_name = "Table"

    suffix = f"_tbl{table_number}"
    max_base_length = MAX_SHEET_NAME_LENGTH - len(suffix)
    sheet_name = f"{base_name[:max_base_length]}{suffix}"

    original_name = sheet_name
    counter = 1

    while sheet_name in used_names:
        counter_suffix = f"_{counter}"
        max_length = MAX_SHEET_NAME_LENGTH - len(counter_suffix)
        sheet_name = f"{original_name[:max_length]}{counter_suffix}"
        counter += 1

    used_names.add(sheet_name)
    return sheet_name


def word_table_to_dataframe(table) -> pd.DataFrame:
    """Convert a Word table to a DataFrame and skip fully empty rows."""
    rows = []

    for row in table.rows:
        cleaned_row = [clean_text(cell.text) for cell in row.cells]

        if any(cleaned_row):
            rows.append(cleaned_row)

    if not rows:
        return pd.DataFrame()

    max_columns = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_columns - len(row)) for row in rows]

    return pd.DataFrame(normalized_rows)


def excel_sheet_to_rows(sheet) -> list[list[str]]:
    """Read an Excel sheet as clean rows and remove empty rows and columns."""
    rows = []

    for excel_row in sheet.iter_rows(values_only=True):
        cleaned_row = [clean_text(value) for value in excel_row]

        if any(cleaned_row):
            rows.append(cleaned_row)

    if not rows:
        return []

    non_empty_columns = [
        column_index
        for column_index in range(max(len(row) for row in rows))
        if any(column_index < len(row) and row[column_index] for row in rows)
    ]

    if not non_empty_columns:
        return []

    return [
        [row[column_index] if column_index < len(row) else "" for column_index in non_empty_columns]
        for row in rows
    ]


def convert_docx_tables_to_excel(input_folder: Path, output_excel_file: Path, log: Callable[[str], None]) -> int:
    """Extract all tables from all .docx files in a folder into one .xlsx file."""
    docx_files = sorted(
        file for file in input_folder.glob("*.docx") if not file.name.startswith("~$")
    )

    if not docx_files:
        raise RuntimeError("В выбранной папке нет файлов .docx.")

    output_excel_file.parent.mkdir(parents=True, exist_ok=True)
    used_sheet_names = set()
    written_tables = 0

    with pd.ExcelWriter(output_excel_file, engine="openpyxl") as writer:
        for docx_file in docx_files:
            try:
                document = Document(docx_file)
            except Exception as error:
                log(f"Ошибка: не удалось открыть '{docx_file.name}': {error}")
                continue

            if not document.tables:
                log(f"Файл '{docx_file.name}': таблицы не найдены.")
                continue

            for table_number, table in enumerate(document.tables, start=1):
                log(f"Word → Excel: {docx_file.name}, таблица №{table_number}")
                dataframe = word_table_to_dataframe(table)

                if dataframe.empty:
                    log(f"Пропуск пустой таблицы: {docx_file.name}, таблица №{table_number}")
                    continue

                sheet_name = make_unique_sheet_name(
                    file_name=docx_file.name,
                    table_number=table_number,
                    used_names=used_sheet_names,
                )

                dataframe.to_excel(
                    writer,
                    sheet_name=sheet_name,
                    index=False,
                    header=False,
                )
                written_tables += 1

        if written_tables == 0:
            pd.DataFrame([["Таблицы не найдены или все таблицы пустые."]]).to_excel(
                writer,
                sheet_name="Info",
                index=False,
                header=False,
            )

    return written_tables


def convert_excel_to_docx(input_excel_file: Path, output_docx_file: Path, log: Callable[[str], None]) -> int:
    """Convert every non-empty sheet from one .xlsx file into Word tables."""
    if input_excel_file.name.startswith("~$") or input_excel_file.suffix.lower() != ".xlsx":
        raise RuntimeError("Выберите обычный файл Excel с расширением .xlsx.")

    try:
        workbook = load_workbook(input_excel_file, data_only=True, read_only=True)
    except Exception as error:
        raise RuntimeError(f"Не удалось открыть Excel-файл: {error}") from error

    output_docx_file.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Таблицы из Excel", level=1)
    document.add_paragraph(f"Источник: {input_excel_file.name}")

    written_tables = 0

    for sheet in workbook.worksheets:
        rows = excel_sheet_to_rows(sheet)

        if not rows:
            log(f"Лист '{sheet.title}': нет данных, пропуск.")
            continue

        log(f"Excel → Word: лист '{sheet.title}', строк: {len(rows)}")
        document.add_heading(sheet.title, level=2)
        table = document.add_table(rows=len(rows), cols=len(rows[0]))

        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        for row_index, row in enumerate(rows):
            for column_index, cell_text in enumerate(row):
                table.cell(row_index, column_index).text = cell_text

        written_tables += 1
        document.add_paragraph()

    workbook.close()

    if written_tables == 0:
        document.add_paragraph("В Excel-файле не найдено непустых листов.")

    try:
        document.save(output_docx_file)
    except Exception as error:
        raise RuntimeError(f"Не удалось сохранить Word-файл: {error}") from error

    return written_tables


class ConverterApp(tk.Tk):
    def __init__(self) -> None:
        if TKINTER_ERROR is not None:
            raise RuntimeError(
                "Не удалось запустить интерфейс: в Python не установлен tkinter. "
                "На Windows установите Python с сайта python.org."
            ) from TKINTER_ERROR

        super().__init__()
        self.title(APP_TITLE)
        self.geometry("900x700")
        self.minsize(800, 600)

        try:
            self.iconbitmap(default="")
        except Exception:
            pass

        self.word_folder_var = tk.StringVar(value=str(Path.home() / "Documents"))
        self.word_output_var = tk.StringVar(value=str(Path.home() / "Documents" / "tables.xlsx"))
        self.excel_file_var = tk.StringVar(value=str(Path.home() / "Documents" / "tables.xlsx"))
        self.excel_output_var = tk.StringVar(value=str(Path.home() / "Documents" / "tables.docx"))
        self.message_queue: Queue[tuple[str, object]] = Queue()
        self.worker_thread: threading.Thread | None = None

        self._build_ui()
        self._poll_queue()

    def _build_ui(self) -> None:
        self._configure_style()
        self.columnconfigure(0, weight=1)
        self.rowconfigure(2, weight=1)

        self._build_menu()

        header_frame = ttk.Frame(self, padding=(16, 16, 16, 8))
        header_frame.grid(row=0, column=0, sticky="ew")
        header_frame.columnconfigure(0, weight=1)

        ttk.Label(
            header_frame,
            text="Magic Button",
            style="Title.TLabel",
        ).grid(row=0, column=0, sticky="w")
        ttk.Label(
            header_frame,
            text="Конвертер таблиц между Word и Excel",
            style="Subtitle.TLabel",
        ).grid(row=1, column=0, pady=(2, 0), sticky="w")

        self.notebook = ttk.Notebook(self)
        self.notebook.grid(row=1, column=0, padx=16, pady=(0, 10), sticky="ew")
        self.notebook.bind("<<NotebookTabChanged>>", self._update_start_button_text)

        self.word_tab = ttk.Frame(self.notebook, padding=12)
        self.excel_tab = ttk.Frame(self.notebook, padding=12)
        self.info_tab = ttk.Frame(self.notebook, padding=16)

        self.notebook.add(self.word_tab, text="📄 Word → Excel")
        self.notebook.add(self.excel_tab, text="📊 Excel → Word")
        self.notebook.add(self.info_tab, text="ℹ️ Справка")

        self._build_word_tab()
        self._build_excel_tab()
        self._build_info_tab()

        work_frame = ttk.Frame(self, padding=(16, 0, 16, 16))
        work_frame.grid(row=2, column=0, sticky="nsew")
        work_frame.columnconfigure(0, weight=1)
        work_frame.rowconfigure(2, weight=1)

        control_frame = ttk.Frame(work_frame)
        control_frame.grid(row=0, column=0, sticky="ew")
        control_frame.columnconfigure(1, weight=1)

        self.start_button = ttk.Button(
            control_frame,
            text="Начать: Word → Excel",
            command=self.start_conversion,
            style="Accent.TButton",
        )
        self.start_button.grid(row=0, column=0, sticky="w")

        self.progress_bar = ttk.Progressbar(control_frame, mode="indeterminate")
        self.progress_bar.grid(row=0, column=1, padx=(12, 0), sticky="ew")

        ttk.Label(work_frame, text="Журнал обработки").grid(
            row=1, column=0, pady=(14, 4), sticky="w"
        )

        log_frame = ttk.Frame(work_frame, borderwidth=1, relief="solid")
        log_frame.grid(row=2, column=0, sticky="nsew")
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)

        self.log_text = tk.Text(
            log_frame,
            height=12,
            wrap="word",
            state="disabled",
            font=("Consolas", 10),
            borderwidth=0,
            padx=8,
            pady=8,
        )
        self.log_text.grid(row=0, column=0, sticky="nsew")

        scrollbar = ttk.Scrollbar(
            log_frame, orient="vertical", command=self.log_text.yview
        )
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.log_text.configure(yscrollcommand=scrollbar.set)

        self.status_var = tk.StringVar(value="Готово к работе")
        status_bar = ttk.Label(
            self,
            textvariable=self.status_var,
            relief="sunken",
            anchor="w",
            padding=(8, 3),
        )
        status_bar.grid(row=3, column=0, sticky="ew")

    def _configure_style(self) -> None:
        self.style = ttk.Style(self)

        if "vista" in self.style.theme_names():
            self.style.theme_use("vista")
        elif "clam" in self.style.theme_names():
            self.style.theme_use("clam")

        self.style.configure(
            "Title.TLabel", font=("Segoe UI", 20, "bold"), foreground="#0078D4"
        )
        self.style.configure(
            "Subtitle.TLabel", font=("Segoe UI", 11), foreground="#606060"
        )
        self.style.configure("Accent.TButton", font=("Segoe UI", 11, "bold"))

    def _build_menu(self) -> None:
        menu_bar = tk.Menu(self)

        file_menu = tk.Menu(menu_bar, tearoff=0)
        file_menu.add_command(label="Выход", command=self.destroy)
        menu_bar.add_cascade(label="Файл", menu=file_menu)

        help_menu = tk.Menu(menu_bar, tearoff=0)
        help_menu.add_command(label="О программе", command=self.show_about)
        menu_bar.add_cascade(label="Справка", menu=help_menu)

        self.config(menu=menu_bar)

    def _build_word_tab(self) -> None:
        self.word_tab.columnconfigure(1, weight=1)

        ttk.Label(
            self.word_tab,
            text="Папка с Word-файлами .docx:",
            font=("Segoe UI", 10, "bold"),
        ).grid(
            row=0,
            column=0,
            padx=(0, 8),
            pady=6,
            sticky="w",
        )
        ttk.Entry(self.word_tab, textvariable=self.word_folder_var).grid(
            row=0, column=1, pady=6, sticky="ew"
        )
        ttk.Button(self.word_tab, text="Выбрать", command=self.choose_word_folder).grid(
            row=0,
            column=2,
            padx=(8, 0),
            pady=6,
        )

        ttk.Label(
            self.word_tab,
            text="Куда сохранить Excel .xlsx:",
            font=("Segoe UI", 10, "bold"),
        ).grid(
            row=1,
            column=0,
            padx=(0, 8),
            pady=6,
            sticky="w",
        )
        ttk.Entry(self.word_tab, textvariable=self.word_output_var).grid(
            row=1, column=1, pady=6, sticky="ew"
        )
        ttk.Button(
            self.word_tab, text="Сохранить как", command=self.choose_word_output
        ).grid(
            row=1,
            column=2,
            padx=(8, 0),
            pady=6,
        )

    def _build_excel_tab(self) -> None:
        self.excel_tab.columnconfigure(1, weight=1)

        ttk.Label(
            self.excel_tab, text="Excel-файл .xlsx:", font=("Segoe UI", 10, "bold")
        ).grid(
            row=0,
            column=0,
            padx=(0, 8),
            pady=6,
            sticky="w",
        )
        ttk.Entry(self.excel_tab, textvariable=self.excel_file_var).grid(
            row=0, column=1, pady=6, sticky="ew"
        )
        ttk.Button(self.excel_tab, text="Выбрать", command=self.choose_excel_file).grid(
            row=0,
            column=2,
            padx=(8, 0),
            pady=6,
        )

        ttk.Label(
            self.excel_tab,
            text="Куда сохранить Word .docx:",
            font=("Segoe UI", 10, "bold"),
        ).grid(
            row=1,
            column=0,
            padx=(0, 8),
            pady=6,
            sticky="w",
        )
        ttk.Entry(self.excel_tab, textvariable=self.excel_output_var).grid(
            row=1, column=1, pady=6, sticky="ew"
        )
        ttk.Button(
            self.excel_tab, text="Сохранить как", command=self.choose_excel_output
        ).grid(
            row=1,
            column=2,
            padx=(8, 0),
            pady=6,
        )

    def _build_info_tab(self) -> None:
        """Build information/help tab."""
        self.info_tab.columnconfigure(0, weight=1)
        self.info_tab.rowconfigure(1, weight=1)

        ttk.Label(
            self.info_tab,
            text="Как пользоваться Magic Button",
            font=("Segoe UI", 14, "bold"),
            foreground="#0078D4",
        ).pack(anchor="w", pady=(0, 12))

        text_frame = ttk.Frame(self.info_tab)
        text_frame.pack(fill="both", expand=True)
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)

        info_text = tk.Text(
            text_frame,
            wrap="word",
            font=("Segoe UI", 10),
            height=20,
            state="normal",
            relief="solid",
            borderwidth=1,
        )
        info_text.grid(row=0, column=0, sticky="nsew")

        scrollbar = ttk.Scrollbar(
            text_frame, orient="vertical", command=info_text.yview
        )
        scrollbar.grid(row=0, column=1, sticky="ns")
        info_text.configure(yscrollcommand=scrollbar.set)

        help_text = """РЕЖИМ 1: WORD → EXCEL
Перенос таблиц из Word-файлов в Excel

Что делает:
• Берет ВСЕ файлы .docx из выбранной папки
• Извлекает из них все таблицы
• Каждую таблицу сохраняет на отдельный лист в Excel
• Убирает пустые строки и лишние пробелы

Как использовать:
1. На вкладке "Word → Excel" нажмите "Выбрать"
2. Укажите папку с .docx файлами
3. Нажмите "Сохранить как" и выберите имя для Excel-файла
4. Нажмите кнопку "Начать"
5. Дождитесь сообщения "Готово"

═══════════════════════════════════════

РЕЖИМ 2: EXCEL → WORD
Перенос листов из Excel в Word-таблицы

Что делает:
• Берет выбранный .xlsx файл
• Все непустые листы преобразует в Word-таблицы
• Каждый лист - это отдельная таблица с заголовком
• Убирает пустые столбцы и лишние пробелы

Как использовать:
1. На вкладке "Excel → Word" нажмите "Выбрать"
2. Выберите .xlsx файл
3. Нажмите "Сохранить как" и выберите имя для Word-файла
4. Нажмите кнопку "Начать"
5. Дождитесь сообщения "Готово"

═══════════════════════════════════════

ТРЕБОВАНИЯ:
• Python 3.7+
• Библиотеки: python-docx, pandas, openpyxl

ПРИМЕЧАНИЯ:
• Поддерживаются только .docx и .xlsx
• Текстовые файлы .doc и .xls не поддерживаются
• В случае ошибки смотрите журнал обработки внизу окна
"""

        info_text.insert("1.0", help_text)
        info_text.configure(state="disabled")

    def choose_word_folder(self) -> None:
        folder = filedialog.askdirectory(title="Выберите папку с .docx-файлами")

        if folder:
            self.word_folder_var.set(folder)

    def choose_word_output(self) -> None:
        file_path = filedialog.asksaveasfilename(
            title="Куда сохранить Excel-файл",
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")],
        )

        if file_path:
            self.word_output_var.set(file_path)

    def choose_excel_file(self) -> None:
        file_path = filedialog.askopenfilename(
            title="Выберите Excel-файл",
            filetypes=[("Excel files", "*.xlsx")],
        )

        if file_path:
            self.excel_file_var.set(file_path)

    def choose_excel_output(self) -> None:
        file_path = filedialog.asksaveasfilename(
            title="Куда сохранить Word-файл",
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
        )

        if file_path:
            self.excel_output_var.set(file_path)

    def start_conversion(self) -> None:
        mode = self._current_mode()

        if mode == "word_to_excel":
            input_path = Path(self.word_folder_var.get().strip())
            output_path = Path(self.word_output_var.get().strip())

            if not input_path.exists() or not input_path.is_dir():
                messagebox.showerror(
                    APP_TITLE, "Выберите существующую папку с .docx-файлами."
                )
                return

            if output_path.suffix.lower() != ".xlsx":
                messagebox.showerror(
                    APP_TITLE, "Выходной файл должен иметь расширение .xlsx."
                )
                return
        else:
            input_path = Path(self.excel_file_var.get().strip())
            output_path = Path(self.excel_output_var.get().strip())

            if not input_path.exists() or not input_path.is_file():
                messagebox.showerror(
                    APP_TITLE, "Выберите существующий Excel-файл .xlsx."
                )
                return

            if input_path.suffix.lower() != ".xlsx":
                messagebox.showerror(
                    APP_TITLE, "Входной файл должен иметь расширение .xlsx."
                )
                return

            if output_path.suffix.lower() != ".docx":
                messagebox.showerror(
                    APP_TITLE, "Выходной файл должен иметь расширение .docx."
                )
                return

        self._clear_log()
        self._set_running(True)
        self._append_log("Старт обработки...")

        self.worker_thread = threading.Thread(
            target=self._run_conversion,
            args=(mode, input_path, output_path),
            daemon=True,
        )
        self.worker_thread.start()

    def _run_conversion(self, mode: str, input_path: Path, output_path: Path) -> None:
        try:
            if mode == "word_to_excel":
                result_count = convert_docx_tables_to_excel(
                    input_folder=input_path,
                    output_excel_file=output_path,
                    log=lambda message: self.message_queue.put(("log", message)),
                )
                result_message = f"✓ Готово. Таблиц сохранено в Excel: {result_count}"
            else:
                result_count = convert_excel_to_docx(
                    input_excel_file=input_path,
                    output_docx_file=output_path,
                    log=lambda message: self.message_queue.put(("log", message)),
                )
                result_message = f"✓ Готово. Таблиц сохранено в Word: {result_count}"
        except Exception as error:
            self.message_queue.put(("error", str(error)))
            return

        self.message_queue.put(("done", result_message))

    def _poll_queue(self) -> None:
        while not self.message_queue.empty():
            message_type, payload = self.message_queue.get()

            if message_type == "log":
                self._append_log(str(payload))
            elif message_type == "error":
                self._set_running(False)
                self._append_log(f"✗ Ошибка: {payload}")
                self.status_var.set("Ошибка")
                messagebox.showerror(APP_TITLE, str(payload))
            elif message_type == "done":
                self._set_running(False)
                message = str(payload)
                self._append_log(message)
                self.status_var.set(message)
                messagebox.showinfo(APP_TITLE, message)

        self.after(100, self._poll_queue)

    def _append_log(self, message: str) -> None:
        self.log_text.configure(state="normal")
        self.log_text.insert("end", f"{message}\n")
        self.log_text.see("end")
        self.log_text.configure(state="disabled")

    def _clear_log(self) -> None:
        self.log_text.configure(state="normal")
        self.log_text.delete("1.0", "end")
        self.log_text.configure(state="disabled")

    def _set_running(self, is_running: bool) -> None:
        if is_running:
            self.start_button.configure(state="disabled")
            self.progress_bar.start(10)
            self.status_var.set("Идет обработка...")
        else:
            self.progress_bar.stop()
            self.start_button.configure(state="normal")
            self._update_start_button_text()

    def _current_mode(self) -> str:
        selected_tab = self.notebook.select()

        if selected_tab == str(self.excel_tab):
            return "excel_to_word"

        return "word_to_excel"

    def _update_start_button_text(self, _event: object | None = None) -> None:
        if not hasattr(self, "start_button"):
            return

        mode = self._current_mode()
        selected_index = self.notebook.index(self.notebook.select())

        if selected_index >= 2:
            self.start_button.configure(state="disabled")
            self.start_button.configure(text="Выберите вкладку")
        else:
            self.start_button.configure(state="normal")
            if mode == "excel_to_word":
                self.start_button.configure(text="Начать: Excel → Word")
            else:
                self.start_button.configure(text="Начать: Word → Excel")

    def show_about(self) -> None:
        messagebox.showinfo(
            APP_TITLE,
            "Magic Button v2.0\n\n"
            "Простая программа для переноса таблиц между Word и Excel.\n\n"
            "Возможности:\n"
            "✓ Word → Excel: все таблицы на отдельные листы\n"
            "✓ Excel → Word: все листы в отдельные таблицы\n"
            "✓ Автоматическая очистка от пустых строк\n"
            "✓ Сохранение форматирования\n\n"
            "Поддерживаемые форматы:\n"
            ".docx (Word) и .xlsx (Excel)\n\n"
            "Требования:\n"
            "Python 3.7+\n"
            "Библиотеки: python-docx, pandas, openpyxl",
        )


if __name__ == "__main__":
    if TKINTER_ERROR is not None:
        raise SystemExit(
            "Не удалось запустить интерфейс: в Python не установлен tkinter. "
            "На Windows установите Python с сайта python.org."
        )

    app = ConverterApp()
    app.mainloop()
